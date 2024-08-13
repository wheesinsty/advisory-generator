import playwright
from playwright.sync_api import sync_playwright
import urllib.request 
from docx import Document
from docx.shared import Inches
from PIL import ImageOps
from PIL import Image
import pandas as pd
import datetime
import os
from PIL import Image, ImageChops
import time

# open the excel sheet
df = pd.read_excel("test.xlsx")

# empty error columns
def emptyErrorColumns():
    for col in df.columns:
        if "Error" in col:
            df[col] = " "
    df.to_excel("test.xlsx", index = False)

# reports start time to excel sheet
def reportStartTime():
    df["Start time"] = df["Start time"].astype(str)
    df.loc[row, "Start time"] = str(datetime.datetime.now().strftime("%H:%M:%S"))
    start_time = time.time()
    df.to_excel("test.xlsx", index = False)

def reportError(msg):
    # report error to excel sheet
    df.loc[row, "Error"] = df.loc[row, "Error"] + "\n " + msg
    df.to_excel("test.xlsx", index = False)
            
# go to customer profile
def goToTenant(df, row): # -> bool:
    # fill customer name into searchbar
    page.locator("#customer-search-box").get_by_placeholder("Search").fill(df.iloc[row]["Domain"])
    page.wait_for_timeout(5000)
    
    # go to customer profile
    page.get_by_role("radio", name = "Select row", checked = False, disabled = False).check()
    
    # go to service management
    page.get_by_text("Service management").click()
    page.wait_for_timeout(5000)
    
    # go to "Microsoft 365"
    try:
        page.locator('//*[@id="MicrosoftOffice"]').click()
        page.wait_for_timeout(5000)
    except:
        reportError("No admin permissions")
        return True

    return False

# fill in the first table
def fillInTable1(df, row):
    doc.tables[0].cell(0, 1).paragraphs[0].add_run(df.iloc[row]["Username"])
    doc.tables[0].cell(1, 1).paragraphs[0].add_run(df.iloc[row]["Domain"])
    doc.tables[0].cell(2, 1).paragraphs[0].add_run(datetime.datetime.now().strftime("%m-%Y"))

# change settings to display usernames in reports >> settings
def displayUsernames(df, row) -> bool:
    # go to reports settings
    page.goto("https://admin.microsoft.com/Adminportal/Home#/Settings/Services/:/Settings/L1/Reports")
 
    page.wait_for_timeout(5000)
    try:       
        # check the "display usernames" checkbox
        page.locator("//html/body/div[4]/div/div/div/div/div[2]/div[2]/div/div[4]/div[4]/label/div/i").check()
    except:
        reportError("No permission to change display usernames")
        return True
    
    page.wait_for_timeout(1000)
    
    try:
        page.get_by_role("button", name = "Save")
    except:
        reportError("Failed to save \"change display usernames\"")
        return True
    
    return False

# take email screenshot
def takeScreenshot(links):
    
    for item in links:
        # go to mailbox usage
        page.goto(links[item])
        page.wait_for_timeout(5000)
        try:
            page.get_by_label("Storage used (MB)").scroll_into_view_if_needed()
        except:
            pass
        
        # sort storage used in descending order     
        for i in range(2):
            try:
                page.wait_for_timeout(1000)
                page.get_by_text("Storage used (MB)").click()
                page.wait_for_timeout(1000)
                page.get_by_label("Sort").click()
            except:
                reportError("Failed to sort " + item + " by storage used")
                return True
                
        # minimise Help & Support button
        if page.get_by_role("button", name = "Minimize Button").is_visible():
            page.get_by_role("button", name = "Minimize Button").click()
            
        page.wait_for_timeout(5000)
        try:
            page.get_by_role("treegrid", name="Usage").screenshot(path = "top" + item + "UsersScreenshot.png")
        except:
            reportError("Failed to screenshot top " + item + " usage")
            return True
        
    return False

# download licenses excel sheet
def downloadLicensesExcel():   
    
    # go to licenses webpage
    page.goto("https://admin.microsoft.com/Adminportal/Home#/licenses")
    page.wait_for_timeout(10000)
    
    # download path of the excel sheet
    download_path = filepath + "/" + df.iloc[row]["Username"] + "_" + str(datetime.datetime.now().strftime("%d-%m-%Y_%H_%M_%S")) + ".csv"
    with page.expect_download() as download_info:
        
        # click the export button
        page.get_by_role("menuitem", name = "Export").click()
    download = download_info.value
    download.save_as(download_path)

# trim excess white parts of the screenshots
def trim(im): 
    #  The background color of the new image is set to the color of the first pixel in im (i.e., im.getpixel((0,0))).
    bg = Image.new(im.mode, im.size, im.getpixel((0,0)))
    diff = ImageChops.difference(im, bg)
    diff = ImageChops.add(diff, diff, 2.0, -100)
    bbox = diff.getbbox()
    if bbox:
        return im.crop(bbox)

# trim excess white parts of the screenshots
def cropImages():   
    im = Image.open("topEmailUsersScreenshot.png")
    trim(im).save("topEmailUsersScreenshot.png")
    im = Image.open("topSharepointUsersScreenshot.png")
    trim(im).save("topSharepointUsersScreenshot.png")
    im = Image.open("topOnedriveUsersScreenshot.png")
    trim(im).save("topOnedriveUsersScreenshot.png")

# add grey border to the screenshots
def add_border(input_image, output_image):
    img = Image.open(input_image)
    bimg = ImageOps.expand(img, border = 20, fill = (500, 500, 500))
    bimg.save(output_image)
    img = Image.open(input_image)
    bimg = ImageOps.expand(img, border = 2, fill = (165, 165, 165))
    bimg.save(output_image)

# add border to screenshots
def addBorder():
    add_border("topEmailUsersScreenshot.png", output_image="topEmailUsersScreenshot.png")
    add_border("topSharepointUsersScreenshot.png", output_image="topSharepointUsersScreenshot.png")
    add_border("topOnedriveUsersScreenshot.png", output_image="topOnedriveUsersScreenshot.png")

# paste screenshots into advisory report
def pasteScreenshots():
    doc.tables[1].cell(0,0).paragraphs[0].add_run().add_picture("topEmailUsersScreenshot.png", width = Inches(5.51))
    doc.tables[1].cell(0,0).paragraphs[3].add_run().add_picture("topSharepointUsersScreenshot.png", width = Inches(5.51))
    doc.tables[1].cell(0,0).paragraphs[7].add_run().add_picture("topOnedriveUsersScreenshot.png", width = Inches(5.51))

# reports end time to excel sheet
def reportEndTime():
    df["End time"] = df["End time"].astype(str)
    df.loc[row, "End time"] = str(datetime.datetime.now().strftime("%H:%M:%S"))
    df.to_excel("test.xlsx", index = False)
    
with sync_playwright() as p:
    
    print("Program start time: " + str(datetime.datetime.now().strftime("%H:%M:%S")))
    browser = p.chromium.connect_over_cdp("http://localhost:9222")
    default_context = browser.contexts[0]
    page = default_context.pages[0]
    emptyErrorColumns()
    links = {"Email": "https://admin.microsoft.com/Adminportal/Home#/reportsUsage/MailboxUsage", "Sharepoint": "https://admin.microsoft.com/Adminportal/Home#/reportsUsage/SharePointSiteUsageV1", "Onedrive": "https://admin.microsoft.com/Adminportal/Home#/reportsUsage/OneDriveSiteUsage"}
    
    for row in range(len(df["Domain"])):
        if df.iloc[row]["Success"] == "Yes" or  df.iloc[row]["Error"] == "No permission to change display usernames" or df.iloc[row]["Error"] == "Failed to save \"change display usernames\"" or df.iloc[row]["Error"] == "No admin permissions": continue
        
        terminateLoop = False
        reportStartTime()
        
        # open the advisory template
        doc = Document('Advisory_OnePageTemplate.docx')
        
        page.goto("https://partner.microsoft.com/dashboard/v2/customers/list")
        
        # navigate to the customer's admin centre
        if goToTenant(df, row): continue
        if displayUsernames(df, row) :continue
        if takeScreenshot(links): continue
        # make a new folder to hold the excel sheet and advisory report of the company
        filepath = "C:/Users/" + os.getenv('USERNAME') + "/Superhub Limited/Cloud Operation Team - CAMP-012 TeammateOnboard/Customer_Report/" + df.iloc[row]["Username"] + "_" + str(datetime.datetime.now().strftime("%d-%m-%Y_%H_%M_%S"))
        os.makedirs(filepath, exist_ok=True)
        fillInTable1(df, row)
        downloadLicensesExcel()
        cropImages()
        addBorder()
        pasteScreenshots()
        
        # save the advisory report into the new folder
        filename = filepath + "/" + df.iloc[row]["Username"] + "_" + str(datetime.datetime.now().strftime("%d-%m-%Y_%H_%M_%S")) + '_Advisory_OnePage.docx'
        try:
            doc.save(filename)
        except:
            reportError("Failed to save file")
            continue
        
        reportEndTime()
        df.loc[row, "Success"] = "Yes"
        df.to_excel("test.xlsx", index = False)

print("Program end time: " + str(datetime.datetime.now().strftime("%H:%M:%S")))       
browser.close()

