"""
automate.py

this module automates the following tasks:
1. Generates of advisory reports for clients
2. Downloads licenses excel sheet for clients
"""
import datetime
import os
import logging
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches
import pandas as pd
from PIL import Image, ImageChops, ImageOps
from docx2pdf import convert

# open the excel sheet
while True:
    input_filepath = input("Please enter the filepath to the users file including the file format (.csv or .xlsx): ")
    if ".csv" in input_filepath:
        try:
            df = pd.read_csv(input_filepath)
        except:
            print("Incorrect filename or format. Please double check the filepath and filetype.")
            continue
    elif ".xlsx" in input_filepath:
        try:
            df = pd.read_excel(input_filepath)
        except:
            print("Incorrect filename or format. Please double check the filepath and filetype.")
            continue
    break

# enter the output filepath
while True:
    output_filepath = input("Please enter the output filepath for the generated reports: ")
    if os.path.exists(output_filepath):
        break

# reset the excel sheet if needed
while True:
    reset = input("Do you want to reset the excel sheet? (yes/no): ")
    if reset.lower() != "yes" and reset.lower() != "no":
        continue
    if reset.lower() == "yes":
        df["Error"] = ""
        df["Success"] = ""
        df["Start time"] = ""
        df["End time"] = ""
        # df.to_excel(input_filepath, index = False)
    break

# reports start time to excel sheet
def reportStartTime():
    if "Start time" not in df.columns:
        df["Start time"] = ""
    df["Start time"] = df["Start time"].astype(str)
    df.loc[row, "Start time"] = str(datetime.datetime.now().strftime("%H:%M:%S"))
    # df.to_excel(input_filepath, index = False)

def reportError(msg):
    # report error to excel sheet
    df.loc[row, "Error"] = msg
    return True
    # df.to_excel(input_filepath, index = False)

def checkIfContinue(row):
    if df.loc[row, "Domain Prefix"] == "": 
        return reportError("No domain name")
    
    if "Success" not in df.columns:
        df["Success"] = ""
        return False
    
    if df.loc[row, "Success"] == "Yes":
        return True
    
    if "Error" not in df.columns: 
        df["Error"] = ""
        return False
    
    errors = [
        "No customers found",
        "No permission to change display usernames",
        "Failed to save \"change display usernames\"",
        "No admin permissions",
        "No service management",
        "No domain name",
    ]
    
    for error in errors:
        if error in str(df.loc[row, "Error"]): 
            return True
    
# go to customer profile
def goToTenant(row):
    # fill customer name into searchbar
    if type(df.loc[row, "Domain Prefix"]) != str or df.loc[row, "Domain Prefix"] == "": 
        return reportError("No domain name")
    
    try:
        page.locator("#customer-search-box").get_by_placeholder("Search").fill(df.loc[row, "Domain Prefix"])
    except:
        return reportError("Cannot fill customer name")

    page.wait_for_timeout(5000)
    
    # go to customer profile
    try:
        page.locator("//html/body/he-layout/div[2]/customer-app/div/customer_he-layout/section[2]/customer_he-data-grid/span[2]/a").click()
    except:
        return reportError("No customers found")

    # go to service management
    try:
        page.get_by_text("Service management").click()
    except:
        return reportError("No service management")

    page.wait_for_timeout(5000)
    
    # go to "Microsoft 365"
    try:
        page.locator('//*[@id="MicrosoftOffice"]').click()
        page.wait_for_timeout(5000)
    except:
        return reportError("No admin permissions")

    page1 = default_context.pages[1]
    page1.close()
    return False

# fill in the first table
def fillInTable1(row):
    doc.tables[0].cell(0, 1).paragraphs[0].add_run(df.loc[row, "Account Description"])
    doc.tables[0].cell(1, 1).paragraphs[0].add_run(df.loc[row, "Domain Prefix"])
    doc.tables[0].cell(2, 1).paragraphs[0].add_run(datetime.datetime.now().strftime("%m-%Y"))
    
# change settings to display usernames in reports >> settings
def displayUsernames():
    # go to reports settings
    page.goto("https://admin.microsoft.com/Adminportal/Home#/Settings/Services/:/Settings/L1/Reports")
 
    page.wait_for_timeout(5000)
    try:       
        # check the "display usernames" checkbox
        page.locator("//html/body/div[4]/div/div/div/div/div[2]/div[2]/div/div[3]/div[4]/label/div/i").check()
    except:
        return reportError("No permission to change display usernames")
    
    page.wait_for_timeout(1000)
    
    try:
        page.get_by_role("button", name = "Save")
    except:
        return reportError("Failed to save \"change display usernames\"")
    
    return False

# take email screenshot
def takeScreenshot(links):
    
    for item in links:
        # go to mailbox usage
        page.goto(links[item])
        page.wait_for_timeout(5000)
        try:
            page.get_by_label("Storage used (MB)").scroll_into_view_if_needed()
        except:
            pass
        
        # sort storage used in descending order     
        for i in range(2):
            try:
                page.wait_for_timeout(1000)
                page.get_by_text("Storage used (MB)").click()
                page.wait_for_timeout(1000)
                page.get_by_text("Sort").click()
            except:
                return reportError("Failed to sort " + item + " by storage used")
                
        # minimise Help & Support button
        if page.get_by_role("button", name="Minimize Button").is_visible():
            page.get_by_role("button", name="Minimize Button").click()
            
        page.wait_for_timeout(5000)
        try:
            page.get_by_role("treegrid", name="Usage").screenshot(path="top"+item+"UsersScreenshot.png")
        except:
            return reportError("Failed to screenshot top " + item + " usage")
        
    return False

# download licenses excel sheet
def downloadLicensesExcel(filepath):   
    
    # go to licenses webpage
    page.goto("https://admin.microsoft.com/Adminportal/Home#/licenses")
    page.wait_for_timeout(40000)
    try:
        with page.expect_download() as download_info:
            
            # click the export button
            try:
                page.get_by_role("menuitem", name="Export").click()
            except:
                return reportError("Failed to click export, possibly because page took too long to load")

    except:
        return reportError("Failed to download licenses excel sheet, possibly because page took too long to load")

    # download path of the excel sheet
    download_path = filepath + "/" + df.loc[row, "Account Description"] + "_" + str(datetime.datetime.now().strftime("%d-%m-%Y_%H_%M_%S")) + ".csv"
    download = download_info.value
    try:
        download.save_as(download_path)
    except:
        return reportError("Failed to save licenses excel sheet to the correct directory, please check that the Account Description has no special characters")

    return False

# trim excess white parts of the screenshots
def trim(im): 
    #  The background color of the new image is set to the color of the first pixel in im (i.e., im.getpixel((0,0))).
    bg = Image.new(im.mode, im.size, im.getpixel((0,0)))
    diff = ImageChops.difference(im, bg)
    diff = ImageChops.add(diff, diff, 2.0, -100)
    bbox = diff.getbbox()
    if bbox:
        return im.crop(bbox)

# trim excess white parts of the screenshots
def cropImages():   
    im = Image.open("topEmailUsersScreenshot.png")
    trim(im).save("topEmailUsersScreenshot.png")
    im = Image.open("topSharepointUsersScreenshot.png")
    trim(im).save("topSharepointUsersScreenshot.png")
    im = Image.open("topOnedriveUsersScreenshot.png")
    trim(im).save("topOnedriveUsersScreenshot.png")

# add grey border to the screenshots
def add_border(input_image, output_image):
    img = Image.open(input_image)
    bimg = ImageOps.expand(img, border = 20, fill = (500, 500, 500))
    bimg.save(output_image)
    img = Image.open(input_image)
    bimg = ImageOps.expand(img, border = 2, fill = (165, 165, 165))
    bimg.save(output_image)

# add border to screenshots
def addBorder():
    add_border("topEmailUsersScreenshot.png", output_image="topEmailUsersScreenshot.png")
    add_border("topSharepointUsersScreenshot.png", output_image="topSharepointUsersScreenshot.png")
    add_border("topOnedriveUsersScreenshot.png", output_image="topOnedriveUsersScreenshot.png")

# paste screenshots into advisory report
def pasteScreenshots():
    doc.tables[1].cell(0,0).paragraphs[0].add_run().add_picture("topEmailUsersScreenshot.png", width = Inches(5.51))
    doc.tables[1].cell(0,0).paragraphs[3].add_run().add_picture("topSharepointUsersScreenshot.png", width = Inches(5.51))
    doc.tables[1].cell(0,0).paragraphs[7].add_run().add_picture("topOnedriveUsersScreenshot.png", width = Inches(5.51))

# save excel sheet
def saveExcel(row, input_filepath):
    if ".csv" in input_filepath:
        df.to_csv(input_filepath, index = False)
    else:
        df.to_excel(input_filepath, index = False)
    
# print program execution time
def programExecutionTime(
    start_time, end_time):
    start_dt = datetime.datetime.strptime(start_time, "%H:%M:%S")
    end_dt = datetime.datetime.strptime(end_time, "%H:%M:%S")

    # Calculate the difference
    time_difference = end_dt - start_dt

    # Extract hours, minutes, and seconds
    hours, remainder = divmod(time_difference.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)

    # Print program execution time
    print("Finished! Program took " + f"{hours} hours, {minutes} minutes, {seconds} seconds")
    
# reports end time to excel sheet
def reportEndTime():
    if "End time" not in df.columns:
        df["End time"] = ""
    df["End time"] = df["End time"].astype(str)
    df.loc[row, "End time"] = str(datetime.datetime.now().strftime("%H:%M:%S"))
    # df.to_excel(input_filepath, index = False)
    
with sync_playwright() as p:
    # display program start time
    start_time = datetime.datetime.now().strftime("%H:%M:%S")
    print("Program start time: " + str(start_time))
    
    # connect to the browser
    browser = p.chromium.connect_over_cdp("http://localhost:9222")
    default_context = browser.contexts[0]
    page = default_context.pages[0]
    
    for row in range(len(df["Domain Prefix"])):
        
        condition = False
        # check if the program should continue with the current client
        if checkIfContinue(row): 
            continue
        
        print("\n Generating report for: " + df.loc[row, "Account Description"])
        print(df)
        reportStartTime()
        
        # open the advisory template
        doc = Document('Advisory_OnePageTemplate.docx')
        
        # go to the partner centre page
        try:
            page.goto("https://partner.microsoft.com/dashboard/v2/customers/list")
        except Exception as e:
            logging.error(f"An error occurred: {e}")
            
        # navigate to the customer's admin centre
        if goToTenant(row): 
            continue
        
        if displayUsernames():
            continue
        
        if takeScreenshot(
                {"Email": "https://admin.microsoft.com/Adminportal/Home#/reportsUsage/MailboxUsage", 
                "Sharepoint": "https://admin.microsoft.com/Adminportal/Home#/reportsUsage/SharePointSiteUsageV1", 
                "Onedrive": "https://admin.microsoft.com/Adminportal/Home#/reportsUsage/OneDriveSiteUsage"}): 
            continue
        
        fillInTable1(row)
        
        # make a new folder to hold the excel sheet and advisory report of the company
        filepath = output_filepath + "/" + df.loc[row, "Account Description"] + "_" + str(datetime.datetime.now().strftime("%d-%m-%Y_%H_%M_%S"))
        os.makedirs(filepath, exist_ok=True)
        
        condition = downloadLicensesExcel(filepath)
        
        # adjust screenshots to fit template
        cropImages()
        addBorder()
        pasteScreenshots()
        
        # save the advisory report into the new folder
        filename = filepath + "/" + df.loc[row, "Account Description"] + "_" + str(datetime.datetime.now().strftime("%d-%m-%Y_%H_%M_%S")) + '_Advisory_OnePage.docx'
        try:
            doc.save(filename)
        except:
            condition = reportError("Failed to save file")
        
        # save the advisory report as pdf
        if condition != True:
            convert(filename, filepath + "/" + df.loc[row, "Account Description"] + "_Advisory_OnePage_" +  str(datetime.datetime.now().strftime("%b%Y")) + ".pdf")
            
        # report end time to excel sheet
        reportEndTime()
        
        # update excel sheet if the report was successful generated
        if condition != True:
            df.loc[row, "Success"] = "Yes"
            df.loc[row, "Error"] = ""
        
        # update excel sheet
        saveExcel(row, input_filepath)

# display program end time and execution time
end_time = datetime.datetime.now().strftime("%H:%M:%S")
print("Program end time: " + str(end_time))
programExecutionTime(start_time, end_time)

# close the browser connection
browser.close()
